import os
import json
import base64
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
import chromadb
from openai import OpenAI
import tiktoken
import uuid

load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
CHROMA_API_KEY = os.getenv('CHROMA_API_KEY')

chroma_client = chromadb.CloudClient(
  api_key=CHROMA_API_KEY,
  tenant='e261f061-5696-47a2-979a-3ca2bb4ee6da',
  database='polydoc'
)

client = OpenAI()

collection = chroma_client.get_or_create_collection(name="translations", configuration={
        "hnsw": {
            "space": "cosine"
        }
    })


def extract_paragraph_runs(element, runs):
    for p in element.paragraphs:
        for r in p.runs:
            if len(r.text) > 0:
                runs.append(r)


def get_non_empty_runs(doc):
    runs = []
    # Paragraphs
    extract_paragraph_runs(doc, runs)

    # Tables
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                extract_paragraph_runs(c, runs)

    # Sections
    for s in doc.sections:
        h = s.header
        extract_paragraph_runs(h, runs)

        f = s.footer
        extract_paragraph_runs(f, runs)

    return runs


def get_text_segments(runs):
    text_segments = []
    for run in runs:
        text_segments.append(run.text)
    return text_segments


class TranslatedTextSegments(object):
    # minimal shape compatible with client.responses.parse usage in original
    def __init__(self, text_segments):
        self.text_segments = text_segments


def translate_text_segments(text_segments, language, client):
    system_prompt = f"""
You are translating a document. You will be given the target language and an array of arrays. Each array represents a paragraph. Individual elements are "runs" (piece of text with the same formatting). You have to translate all paragraphs and return back the same representation using arrays, but translated. Return back only the translated array.

Translate each paragraph to natural {language}, preserving meaning and idiomatic expressions, rather than translating word-for-word. Maintain the same run structure, but aim for fluent, localized language that a {language} speaker would naturally use.
    """
    
    user_prompt = f"""
Target language: {language}

Document representation: {json.dumps(text_segments, ensure_ascii=False)}
    """

    res = client.responses.parse(
        model="gpt-4o",
        input=[
            {"role": "system", "content": system_prompt.strip()},
            {"role": "user", "content": user_prompt.strip()},
        ],
        text_format=TranslatedTextSegments
    )
    
    translated_arr = res.output_parsed.text_segments
    
    return translated_arr


def retrieve_past_translations(text_segments, translations, language_code, client):
    to_be_translated = []
    no_of_reused = 0
    no_of_total = 0
    tokens = 0

    encoding = tiktoken.encoding_for_model("gpt-4o")

    for ts in text_segments:
        query_emb = client.embeddings.create(
            model="text-embedding-3-small",
            input=ts
        ).data[0].embedding

        results = collection.query(
            query_embeddings=[query_emb],
            n_results=1
        )

        # Safety: check for empty results
        if len(results["distances"]) == 0 or len(results["distances"][0]) == 0:
            to_be_translated.append(ts)
            no_of_total += 1
            continue

        distance = results["distances"][0][0]

        if distance < 0.3:
            metadata_list = results.get("metadatas", [[]])[0][0]  # get first result metadata dict
            if metadata_list and language_code in metadata_list:
                translations[ts] = metadata_list[language_code]
                no_of_reused += 1
                tokens += len(encoding.encode(ts))
            else:
                # fallback if metadata doesn't have the language
                to_be_translated.append(ts)
        else:
            to_be_translated.append(ts)

        no_of_total += 1
    
    cost_savings = tokens / 1000000 * 2.5 + tokens / 1000000 * 10

    return {
        "to_be_translated": to_be_translated,
        "no_of_reused": no_of_reused,
        "no_of_total": no_of_total,
        "cost_savings": cost_savings
    }


def translate_document_and_return(doc, language_full, language_code, client):
    runs = get_non_empty_runs(doc)
    text_segments = get_text_segments(runs)

    translations = {}
    retrieved_data = retrieve_past_translations(text_segments, translations, language_code, client)

    to_be_translated = retrieved_data["to_be_translated"]
    translated_text_segments = []
    if to_be_translated:
        translated_text_segments = translate_text_segments(to_be_translated, language_full, client)

    # Save translations into the index
    for idx, translated_text_segment in enumerate(translated_text_segments):
        source_text = to_be_translated[idx]
        translations[source_text] = translated_text_segment

        # Add to Chroma
        unique_id = str(uuid.uuid4())
        emb = client.embeddings.create(
            model="text-embedding-3-small",
            input=to_be_translated[idx]
        )
        vector = emb.data[0].embedding
        collection.add(
            ids=[unique_id],
            documents=[source_text],
            embeddings=[vector],
            metadatas=[{
                language_code: translated_text_segment
            }]
        )
    
    for run in runs:
        run.text = translations.get(run.text, run.text)

    # Save translated docx to bytes
    output_buffer = BytesIO()
    doc.save(output_buffer)
    doc_bytes = output_buffer.getvalue()

    return doc_bytes, {
        "total_segments": retrieved_data["no_of_total"],
        "reused_segments": retrieved_data["no_of_reused"],
        "cost_savings": retrieved_data["cost_savings"],
    }


def lambda_handler(event, context):
    try:
        encoded_doc = event.get('file')
        language = event.get('language')
        language_code = event.get('languageCode')

        if not all([encoded_doc, language, language_code]):
            return {"statusCode": 400, "error": "Missing required parameters"}

        try:
            document_data = base64.b64decode(encoded_doc)
            input_buffer = BytesIO(document_data)
        except Exception as e:
            return {"statusCode": 400, "error": f"Invalid base64 document: {str(e)}"}

        try:
            doc = Document(input_buffer)
        except Exception as e:
            return {"statusCode": 400, "error": f"Failed to parse DOCX: {str(e)}"}

        doc_bytes, metrics = translate_document_and_return(doc, language, language_code, client)

        translated_b64 = base64.b64encode(doc_bytes).decode('utf-8')

        return {
            "statusCode": 200,
            "translated_file": translated_b64,
            "total_segments": metrics["total_segments"],
            "reused_segments": metrics["reused_segments"],
            "cost_savings": metrics["cost_savings"]
        }

    except Exception as e:
        return {"statusCode": 500, "error": str(e)}
